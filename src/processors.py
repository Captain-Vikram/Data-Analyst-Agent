"""
File processing utilities for the AI Data Analyst Agent.

This module contains the FileProcessor class that handles reading and
processing various file formats including CSV, Excel, PDF, DOCX, and images.
"""

import os
import tempfile
from typing import Dict, Any, Optional
import warnings
warnings.filterwarnings('ignore')

# Data processing libraries
import pandas as pd
import numpy as np


class FileProcessor:
    """
    Handles processing of various file formats.
    
    This class provides methods to read and process different file types,
    extracting data and metadata as appropriate for each format.
    """
    
    def __init__(self):
        """Initialize the FileProcessor."""
        pass
    
    def process_csv(self, file_path: str) -> Dict[str, Any]:
        """
        Process CSV files.
        
        Args:
            file_path: Path to the CSV file
            
        Returns:
            Dictionary containing the processed data and metadata
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin1', 'iso-8859-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                return {'error': 'Could not read CSV file with any encoding'}
            
            # Basic info
            info = {
                'rows': len(df),
                'columns': len(df.columns),
                'memory_usage': df.memory_usage(deep=True).sum(),
                'dtypes': df.dtypes.to_dict(),
                'missing_values': df.isnull().sum().to_dict(),
                'shape': df.shape
            }
            
            return {
                'data': df,
                'info': info,
                'type': 'csv'
            }
            
        except Exception as e:
            return {'error': f'CSV processing failed: {str(e)}'}
    
    def process_excel(self, file_path: str) -> Dict[str, Any]:
        """
        Process Excel files.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Dictionary containing the processed data and metadata
        """
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            if len(excel_file.sheet_names) == 1:
                # Single sheet
                df = pd.read_excel(file_path)
                info = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'memory_usage': df.memory_usage(deep=True).sum(),
                    'dtypes': df.dtypes.to_dict(),
                    'missing_values': df.isnull().sum().to_dict(),
                    'shape': df.shape,
                    'sheets': excel_file.sheet_names
                }
                
                return {
                    'data': df,
                    'info': info,
                    'type': 'excel'
                }
            else:
                # Multiple sheets
                sheets_data = {}
                total_info = {'sheets': excel_file.sheet_names, 'sheet_details': {}}
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    sheets_data[sheet_name] = df
                    
                    total_info['sheet_details'][sheet_name] = {
                        'rows': len(df),
                        'columns': len(df.columns),
                        'shape': df.shape
                    }
                
                # Return first sheet as main data
                main_sheet = excel_file.sheet_names[0]
                return {
                    'data': sheets_data[main_sheet],
                    'info': total_info,
                    'type': 'excel',
                    'all_sheets': sheets_data
                }
                
        except Exception as e:
            return {'error': f'Excel processing failed: {str(e)}'}
    
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Process PDF files.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    text += page.extract_text() + "\n"
                
                # Get metadata
                metadata = pdf_reader.metadata if pdf_reader.metadata else {}
                
                info = {
                    'pages': len(pdf_reader.pages),
                    'word_count': len(text.split()),
                    'character_count': len(text),
                    'metadata': {str(k): str(v) for k, v in metadata.items()} if metadata else {}
                }
                
                return {
                    'text': text,
                    'info': info,
                    'type': 'pdf'
                }
                
        except ImportError:
            return {
                'error': 'PyPDF2 not installed. Install with: pip install PyPDF2',
                'message': 'PDF processing requires PyPDF2 library'
            }
        except Exception as e:
            return {'error': f'PDF processing failed: {str(e)}'}
    
    def process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Process DOCX files.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            table_text = ""
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_text += cell.text + "\t"
                    table_text += "\n"
            
            full_text = text + "\n" + table_text
            
            info = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'word_count': len(full_text.split()),
                'character_count': len(full_text)
            }
            
            return {
                'text': full_text,
                'info': info,
                'type': 'docx'
            }
            
        except ImportError:
            return {
                'error': 'python-docx not installed. Install with: pip install python-docx',
                'message': 'DOCX processing requires python-docx library'
            }
        except Exception as e:
            return {'error': f'DOCX processing failed: {str(e)}'}
    
    def process_image(self, file_path: str) -> Dict[str, Any]:
        """
        Process image files using OCR.
        
        Args:
            file_path: Path to the image file
            
        Returns:
            Dictionary containing extracted text and image metadata
        """
        try:
            from PIL import Image
            import pytesseract
            
            # Open and process image
            image = Image.open(file_path)
            
            # Extract text using OCR
            try:
                text = pytesseract.image_to_string(image)
            except pytesseract.TesseractNotFoundError:
                return {
                    'error': 'Tesseract OCR engine not installed',
                    'message': '''Tesseract OCR engine is required for image text extraction.
                    
Install instructions:
- Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki
- macOS: brew install tesseract
- Linux: sudo apt-get install tesseract-ocr

Then install Python wrapper: pip install pytesseract'''
                }
            
            # Get image info
            info = {
                'size': image.size,
                'mode': image.mode,
                'format': image.format,
                'word_count': len(text.split()) if text.strip() else 0,
                'character_count': len(text) if text else 0
            }
            
            return {
                'text': text,
                'info': info,
                'type': 'image'
            }
            
        except ImportError as e:
            missing_lib = 'PIL/Pillow' if 'PIL' in str(e) else 'pytesseract'
            return {
                'error': f'{missing_lib} not installed',
                'message': f'Image processing requires {missing_lib}. Install with: pip install Pillow pytesseract'
            }
        except Exception as e:
            return {'error': f'Image processing failed: {str(e)}'}
